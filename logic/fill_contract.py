import customtkinter as ctk
from tkinter import simpledialog
from docx import Document
from tkinter import filedialog
import os

def fill_existing_template():
    file_path = filedialog.askopenfilename(
        title="Choose Contract Template",
        filetypes=[("Word Documents", "*.docx")]
    )

    if not file_path:
        return

    doc = Document(file_path)

    # Step 1: Extract placeholders from text (e.g., {{ClientName}})
    placeholders = set()
    for para in doc.paragraphs:
        if "{{" in para.text and "}}" in para.text:
            words = para.text.split()
            for word in words:
                if word.startswith("{{") and word.endswith("}}"):
                    placeholders.add(word.strip("{}"))

    if not placeholders:
        ctk.CTkMessagebox(title="No Placeholders", message="No {{placeholders}} found in template.")
        return

    # Step 2: Ask user to fill values
    values = {}
    for placeholder in placeholders:
        user_input = simpledialog.askstring("Fill Contract", f"Enter value for: {placeholder}")
        values[placeholder] = user_input

    # Step 3: Replace placeholders
    for para in doc.paragraphs:
        for placeholder, value in values.items():
            para.text = para.text.replace(f"{{{{{placeholder}}}}}", value)

    # Step 4: Save filled contract
    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")],
        title="Save Filled Contract"
    )
    if save_path:
        doc.save(save_path)
        print(f"✅ Filled contract saved at: {save_path}")
