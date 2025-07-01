import customtkinter as ctk
from tkinter import filedialog
from docx import Document

def create_contract_from_form():
    form_window = ctk.CTkToplevel()
    form_window.title("🧾 Create New Contract")
    form_window.geometry("500x500")

    # Field labels and entry boxes
    fields = {
        "Client Name": None,
        "Service Provider": None,
        "Service Description": None,
        "Start Date": None,
        "End Date": None,
        "Payment Amount": None
    }

    entries = {}

    for i, (label, _) in enumerate(fields.items()):
        lbl = ctk.CTkLabel(form_window, text=label)
        lbl.pack(pady=(10 if i == 0 else 5, 0))
        entry = ctk.CTkEntry(form_window, width=400)
        entry.pack()
        entries[label] = entry

    def submit():
        # Get values
        values = {label: entry.get() for label, entry in entries.items()}

        # Generate DOCX document
        doc = Document()
        doc.add_heading("Service Agreement", level=1)
        doc.add_paragraph(f"This Service Agreement is made between {values['Client Name']} and {values['Service Provider']}.")
        doc.add_paragraph(f"The service to be provided is: {values['Service Description']}.")
        doc.add_paragraph(f"Start Date: {values['Start Date']} | End Date: {values['End Date']}")
        doc.add_paragraph(f"Payment Terms: {values['Payment Amount']}")

        # Save path
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Save Generated Contract"
        )

        if save_path:
            doc.save(save_path)
            form_window.destroy()
            ctk.CTkMessagebox(title="Saved", message="✅ New contract created successfully!")

    # Submit button
    submit_btn = ctk.CTkButton(form_window, text="Generate Contract", command=submit)
    submit_btn.pack(pady=20)
